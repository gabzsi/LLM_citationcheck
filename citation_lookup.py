"""
━━━━━━━━━━━━━━━━━━
Uses the CrossRef XML API with unstructured_citation — the same engine
that powers apps.crossref.org/SimpleTextQuery.

Step 1: Submit raw citation → CrossRef Simple Text Query → DOI + basic fields
Step 2: Fetch full metadata from REST API using returned DOI → title, publisher

Output: Excel with one row per citation.
"""

import re
import os
import time
import html
import tkinter as tk
from tkinter import filedialog

import requests
import pandas as pd
from docx import Document

# ── Config ────────────────────────────────────────────────────────────────────
EMAIL           = "your_email@example.com"  # required by CrossRef
REQUEST_DELAY   = 0.5
REQUEST_TIMEOUT = 15
DEBUG           = False   # set True to print raw response for first citation
# ─────────────────────────────────────────────────────────────────────────────

_HEADERS = {
    "User-Agent":   f"CitationLookup/1.0 (mailto:{EMAIL})",
    "Content-Type": "application/x-www-form-urlencoded",
}
_debug_done = False


# ══════════════════════════════════════════════════════════════════════════════
#  Step 1: Parse CrossRef Simple Text Query response
# ══════════════════════════════════════════════════════════════════════════════

def _parse_response(text: str, original: str) -> dict:
    """
    Parse CrossRef's pipe-delimited plain text response.

    Format per line:
      ISSN(s)|journal|author|volume|issue|page|year|type|key|DOI
    Example:
      0009-2665,1520-6890|Chemical Reviews|Nitopi|119|12|7610|2019|full_text|q1|10.1021/acs.chemrev.8b00705

    Returns one result dict. Status is 'resolved' or 'Not Found'.
    """
    result = {
        "original_citation": original,
        "doi": "", "title": "", "author": "", "journal": "",
        "publisher": "", "year": "", "volume": "", "page": "",
        "status": "", "extracted_journal": "", "journal_publisher": "",
    }

    text = text.strip()
    if not text:
        result['status'] = "Not Found"
        return result

    line = next((l.strip() for l in text.splitlines() if l.strip()), "")
    if not line:
        result['status'] = "Not Found"
        return result

    parts = line.split('|')
    doi   = parts[-1].strip() if parts else ""

    if not doi.startswith('10.'):
        result['status'] = "Not Found"
        return result

    result['doi']    = doi
    result['status'] = "resolved"

    # Format: ISSN|journal|author|volume|issue|page|year|type|key|DOI
    if len(parts) >= 10:
        result['journal'] = parts[1].strip()
        result['author']  = parts[2].strip()
        result['volume']  = parts[3].strip()
        result['page']    = parts[5].strip()
        result['year']    = parts[6].strip()

    return result


# ══════════════════════════════════════════════════════════════════════════════
#  Step 2: Enrich with full metadata via REST API
# ══════════════════════════════════════════════════════════════════════════════

def _lookup_doi(doi: str) -> dict:
    """
    Fetch full metadata from CrossRef REST API for a known DOI.
    Returns title, full journal name, full author string, year, publisher.
    """
    try:
        r = requests.get(
            f"https://api.crossref.org/works/{doi}",
            headers={"User-Agent": f"CitationLookup/1.0 (mailto:{EMAIL})"},
            timeout=REQUEST_TIMEOUT,
        )
        r.raise_for_status()
        work = r.json().get('message', {})

        def _c(t):
            return html.unescape(str(t)).strip()

        title      = _c((work.get('title') or [''])[0])
        containers = work.get('container-title', [])
        journal    = _c(containers[0]) if containers else ''
        publisher  = _c(work.get('publisher', ''))

        authors = work.get('author', [])
        first   = ''
        if authors:
            first = _c(authors[0].get('family') or authors[0].get('name') or '')
        author_str = f"{first} et al." if len(authors) > 1 else first

        year = ''
        for field in ('published-print', 'published-online', 'issued', 'created'):
            dp = work.get(field, {}).get('date-parts')
            if dp and dp[0] and dp[0][0]:
                year = str(dp[0][0])
                break

        return {
            'full_title':   title,
            'full_journal': journal,
            'full_author':  author_str,
            'full_year':    year,
            'publisher':    publisher,
        }
    except Exception:
        return {
            'full_title': '', 'full_journal': '',
            'full_author': '', 'full_year': '', 'publisher': '',
        }


# ══════════════════════════════════════════════════════════════════════════════
#  Step 3: Compare CrossRef result against original citation
# ══════════════════════════════════════════════════════════════════════════════

def _norm(text: str) -> str:
    return re.sub(r'[^a-z0-9\s]', '', str(text).lower()).strip()


def _word_overlap(a: str, b: str) -> float:
    sa, sb = set(a.split()), set(b.split())
    return len(sa & sb) / len(sa) if sa else 0.0


def _stem_overlap(a: str, b: str) -> float:
    sa = {w[:3] for w in a.split() if len(w) >= 3}
    sb = {w[:3] for w in b.split() if len(w) >= 3}
    return len(sa & sb) / len(sa) if sa else 0.0


def _compare(original: str, found: dict) -> dict:
    """
    Compare CrossRef's returned metadata against the original raw citation.

    Each field is checked for presence in the raw citation string:
      - title:   word overlap between found title and raw citation
      - author:  found first-author last name appears in raw citation
      - journal: stem overlap between found journal and raw citation
      - year:    found year appears in raw citation

    Returns match columns compatible with citation_analysis_v3.py.
    """
    raw = _norm(original)

    # Title — word overlap between CrossRef title and raw citation
    found_title = _norm(found.get('title', ''))
    if not found_title:
        title_match = "❌"
    else:
        fwd = _word_overlap(found_title, raw)
        rev = _word_overlap(raw, found_title)
        title_match = "✅" if max(fwd, rev) >= 0.45 else "❌"

    # Author — first author last name from CrossRef appears in raw citation
    # The pipe response gives e.g. "Nitopi", full response gives "Nitopi et al."
    raw_author = _norm(found.get('author', '').split()[0] if found.get('author') else '')
    if not raw_author or raw_author == 'unknown':
        author_match = "❌"
    else:
        author_match = "✅" if raw_author in raw else "❌"

    # Journal — stem overlap between CrossRef journal and raw citation
    found_journal = _norm(found.get('journal', ''))
    if not found_journal or found_journal in ('unknown journal', 'unknown', ''):
        journal_match = "❌"
    else:
        journal_match = "✅" if _stem_overlap(found_journal, raw) >= 0.40 else "❌"

    # Year — CrossRef year appears somewhere in raw citation
    found_year = str(found.get('year', '')).strip()
    year_match = "✅" if found_year and found_year in original else "❌"

    # DOI — extract DOI from original citation and compare to CrossRef's DOI
    orig_doi_m = re.search(
        r'10\.\d{4,9}/\S+', original, flags=re.IGNORECASE
    )
    orig_doi     = orig_doi_m.group(0).rstrip('.,)') if orig_doi_m else ''
    returned_doi = found.get('doi', '')
    if not orig_doi:
        doi_match = "–"   # no DOI in original to compare
    elif not returned_doi.startswith('10.'):
        doi_match = "–"   # CrossRef found nothing
    elif orig_doi.lower() == returned_doi.lower():
        doi_match = "✅"
    else:
        doi_match = "❌"

    match_score = sum(x == "✅" for x in
                      [title_match, author_match, journal_match, year_match])

    # Status and needs_verification
    if found.get('status') != 'resolved' or title_match == "❌":
        status            = "❌"
        needs_verification = "–"
    elif match_score >= 3:
        status            = "✅"
        needs_verification = "❌"
    else:
        status            = "⚠️"
        needs_verification = "✅"

    return {
        'status':            status,
        'needs_verification': needs_verification,
        'title_match':       title_match,
        'author_match':      author_match,
        'journal_match':     journal_match,
        'year_match':        year_match,
        'doi_match':         doi_match,
        'match_score':       match_score,
    }

def _extract_journal_from_citation(citation: str) -> str:
    """
    Extract a likely journal name from a raw citation string using regex.

    Targets common patterns:
      - APA/MLA: text between article title period and volume/issue/year
        e.g. `. Journal Name, 12(3)` or `. Journal Name. 12,`
      - Italic-implied: text after the last period before a number sequence
      - Fallback: return empty string if no match found
    """
    # APA/MLA: period + space + Title Case words + comma/period before volume/issue
    m = re.search(
        r'\.\s+([A-Z][A-Za-z &:\-]+?),\s*\d',
        citation
    )
    if m:
        return m.group(1).strip()

    # After last period, before a number sequence (volume/year)
    m = re.search(
        r'\.\s+([A-Z][A-Za-z &:\-]+?)\.\s*\d',
        citation
    )
    if m:
        return m.group(1).strip()

    return ""


def _lookup_publisher_by_journal(journal_name: str) -> str:
    """
    Query the CrossRef REST API journals endpoint to find the publisher
    for the given journal name.

    Returns the publisher string from message.items[0].publisher,
    or empty string on failure or no results.
    """
    try:
        time.sleep(REQUEST_DELAY)
        r = requests.get(
            "https://api.crossref.org/journals",
            params={"query": journal_name},
            headers={"User-Agent": f"CitationLookup/1.0 (mailto:{EMAIL})"},
            timeout=REQUEST_TIMEOUT,
        )
        r.raise_for_status()
        data = r.json()
        items = data.get('message', {}).get('items', [])
        if items:
            return items[0].get('publisher', '')
        return ""
    except Exception:
        return ""


def query_crossref(citation: str) -> dict:
    """
    Submit a raw citation string to CrossRef's unstructured_citation matcher
    (same engine as apps.crossref.org/SimpleTextQuery), then enrich the
    result with full metadata via the REST API.

    The DOI is stripped from the citation before querying — AI-generated
    DOIs are often hallucinated and confuse CrossRef's matcher.
    The returned DOI comes entirely from CrossRef's own match.
    """
    global _debug_done

    # Strip any DOI from the citation before querying
    # so CrossRef matches on title/author/journal/year only
    citation_no_doi = re.sub(
        r'(?:doi\s*:\s*|https?://doi\.org/|https?://dx\.doi\.org/)?10\.\d{4,9}/\S+',
        '', citation, flags=re.IGNORECASE
    ).strip().rstrip('.,;')

    xml_query = f"""<?xml version="1.0" encoding="UTF-8"?>
<query_batch version="2.0"
  xmlns="http://www.crossref.org/qschema/2.0"
  xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"
  xsi:schemaLocation="http://www.crossref.org/qschema/2.0
    http://www.crossref.org/qschema/crossref_query_input2.0.xsd">
  <head>
    <email_address>{EMAIL}</email_address>
    <doi_batch_id>batch_1</doi_batch_id>
  </head>
  <body>
    <query key="q1" enable-multiple-hits="false" forward-match="false">
      <unstructured_citation>{html.escape(citation_no_doi)}</unstructured_citation>
    </query>
  </body>
</query_batch>"""

    try:
        response = requests.post(
            f"https://doi.crossref.org/servlet/query?usr={EMAIL}",
            data={"qdata": xml_query},
            headers=_HEADERS,
            timeout=REQUEST_TIMEOUT,
        )

        if DEBUG and not _debug_done:
            _debug_done = True
            print(f"\n--- DEBUG HTTP {response.status_code} ---")
            print(f"Content-Type: {response.headers.get('content-type', '?')}")
            print(f"Response:\n{response.text[:800]}")
            print("--- END DEBUG ---\n")

        response.raise_for_status()
        result = _parse_response(response.text, citation)

        # Step 2: enrich with full metadata if a DOI was found
        if result['doi'].startswith('10.'):
            time.sleep(REQUEST_DELAY)
            full = _lookup_doi(result['doi'])
            result['title']     = full['full_title']
            result['publisher'] = full['publisher']
            if not result['journal']:
                result['journal'] = full['full_journal']
            if not result['author']:
                result['author']  = full['full_author']
            if not result['year']:
                result['year']    = full['full_year']

        # Extract journal from original citation string
        result['extracted_journal'] = _extract_journal_from_citation(citation)
        if result['extracted_journal']:
            result['journal_publisher'] = _lookup_publisher_by_journal(result['extracted_journal'])
        else:
            result['journal_publisher'] = ''

        # Step 3: compare CrossRef result against original citation
        matches = _compare(citation, result)
        result.update(matches)

        # found_citation — human-readable summary of what CrossRef returned
        result['found_citation'] = (
            f"{result.get('author','')} ({result.get('year','')}). "
            f"{result.get('title','')}. "
            f"{result.get('journal','')}. "
            f"(DOI: {result.get('doi','')})"
            if result.get('doi', '').startswith('10.')
            else "Not Found"
        )

        return result

    except requests.exceptions.RequestException as e:
        return {
            "original_citation": citation, "found_citation": "Not Found",
            "doi": "API Error", "title": "", "author": "",
            "journal": "", "publisher": "", "year": "",
            "status": "❌", "needs_verification": "–",
            "title_match": "❌", "author_match": "❌",
            "journal_match": "❌", "year_match": "❌",
            "doi_match": "–", "match_score": 0,
            "extracted_journal": "", "journal_publisher": "",
        }


# ══════════════════════════════════════════════════════════════════════════════
#  Document helpers
# ══════════════════════════════════════════════════════════════════════════════

def clean_citation(text: str) -> str:
    text = html.unescape(text).replace('\ufffd', '').strip()
    text = re.sub(r'[\u200b\u200c\u200d\ufeff\u00ad]', '', text)
    text = re.sub(r'^(?:\[\d+\]|\(\d+\)|\d+[\.\)]\s*|[\u2022\-\*]\s*)', '', text)
    return text.strip()


def extract_citations_from_docx(path: str) -> list[str]:
    doc = Document(path)
    return [c for para in doc.paragraphs
            if (c := clean_citation(para.text.strip()))]


def parse_filename(filename: str) -> tuple[str, str]:
    name  = os.path.splitext(filename)[0]
    parts = name.split('_', 1)
    if len(parts) < 2 or not parts[1].strip():
        print(f"  ⚠️  '{filename}' doesn't follow 'engine_topic.docx' convention.")
        return parts[0], "Unknown"
    return parts[0], parts[1]


# ══════════════════════════════════════════════════════════════════════════════
#  Main
# ══════════════════════════════════════════════════════════════════════════════

def main() -> None:
    root = tk.Tk()
    root.withdraw()

    print("Opening file dialog… select your .docx file(s).")
    file_paths = filedialog.askopenfilenames(
        title="Select Word Document(s)",
        filetypes=[("Word Documents", "*.docx")]
    )
    if not file_paths:
        print("No files selected. Exiting.")
        return

    all_results: list[dict] = []

    for file_path in file_paths:
        filename  = os.path.basename(file_path)
        print(f"\n📄 Processing: {filename}")
        ai_engine, topic = parse_filename(filename)

        try:
            citations = extract_citations_from_docx(file_path)
        except Exception as e:
            print(f"  ❌ Could not read {filename}: {e}")
            continue

        print(f"  Found {len(citations)} citations. Querying CrossRef…")

        for i, citation in enumerate(citations, 1):
            time.sleep(REQUEST_DELAY)
            result = query_crossref(citation)
            result['ai_engine'] = ai_engine
            result['topic']     = topic

            doi_str    = result.get('doi', '')    or "—"
            status_str = result.get('status', '') or "—"
            title_str  = result.get('title', '')  or "—"
            if len(title_str) > 60:
                title_str = title_str[:57] + "…"
            print(f"  [{i:>3}/{len(citations)}]  {status_str:<12}  {doi_str:<35}  {title_str}")

            all_results.append(result)

    if not all_results:
        print("\nNo results to save.")
        return

    output_dir   = os.path.dirname(file_paths[0])
    output_excel = os.path.join(output_dir, "Citation_Lookup_Results.xlsx")

    col_order = [
        'ai_engine', 'topic',
        'status', 'needs_verification',
        'title_match', 'author_match', 'journal_match', 'year_match', 'doi_match',
        'match_score',
        'doi', 'title', 'author', 'journal', 'extracted_journal', 'journal_publisher', 'publisher', 'year',
        'original_citation', 'found_citation',
    ]
    df = pd.DataFrame(all_results)
    for col in col_order:
        if col not in df.columns:
            df[col] = ''
    df[col_order].to_excel(output_excel, index=False)

    total     = len(all_results)
    resolved  = sum(1 for r in all_results if r.get('doi', '').startswith('10.'))
    not_found = sum(1 for r in all_results if r.get('status') == 'Not Found')

    print(f"\n{'━'*60}")
    print(f"  Total citations  : {total}")
    print(f"  DOI resolved     : {resolved}/{total}  ({100*resolved//total if total else 0}%)")
    print(f"  Not found        : {not_found}/{total}")
    print(f"{'━'*60}")
    print(f"  Excel → {output_excel}")
    print(f"{'━'*60}\n")


if __name__ == "__main__":
    main()